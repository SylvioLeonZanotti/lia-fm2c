import hashlib
from pathlib import Path
from docx import Document

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_A = "http://schemas.openxmlformats.org/drawingml/2006/main"

_IMAGE_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
_IMG_SIGS  = (b"\xff\xd8", b"\x89PNG", b"GIF8", b"RIFF", b"BM")


def _is_valid_image(image_bytes: bytes) -> bool:
    return any(image_bytes[:4].startswith(sig) for sig in _IMG_SIGS)


def _image_hash(image_bytes: bytes) -> str:
    return hashlib.md5(image_bytes, usedforsecurity=False).hexdigest()


def _map_image_relationships(doc: Document) -> dict[str, bytes]:
    """Returns {relationship_id: image_bytes} for every image in the document."""
    image_by_rid: dict[str, bytes] = {}
    for rid, rel in doc.part.rels.items():
        if rel.reltype == _IMAGE_REL:
            try:
                image_bytes = rel.target_part.blob
                if _is_valid_image(image_bytes):
                    image_by_rid[rid] = image_bytes
            except Exception:
                pass
    return image_by_rid


def _extract_images(file_path: Path) -> list[bytes]:
    """Extracts all unique images from a .docx in document order."""
    doc          = Document(file_path)
    image_by_rid = _map_image_relationships(doc)

    images: list[bytes]   = []
    seen_hashes: set[str] = set()

    for drawing in doc.element.body.iter(f"{{{_W}}}drawing"):
        for blip in drawing.iter(f"{{{_A}}}blip"):
            rid = blip.get(f"{{{_R}}}embed")
            if rid and rid in image_by_rid:
                image_bytes = image_by_rid[rid]
                h = _image_hash(image_bytes)
                if h not in seen_hashes:
                    seen_hashes.add(h)
                    images.append(image_bytes)

    return images


def _extract_text(file_path: Path) -> str:
    """Extracts paragraph text from a .docx, inserting [IMAGEM N] markers where images appear."""
    doc          = Document(file_path)
    image_by_rid = _map_image_relationships(doc)

    blocks: list[str]         = []
    hash_to_index: dict[str, int] = {}
    image_counter = 0

    for para in doc.paragraphs:
        has_drawing = para._element.find(f".//{{{_W}}}drawing") is not None
        text        = para.text.strip()

        if has_drawing:
            if text:
                blocks.append(text)
            img_index: int | None = None
            for drawing in para._element.iter(f"{{{_W}}}drawing"):
                for blip in drawing.iter(f"{{{_A}}}blip"):
                    rid = blip.get(f"{{{_R}}}embed")
                    if rid and rid in image_by_rid:
                        h = _image_hash(image_by_rid[rid])
                        if h not in hash_to_index:
                            image_counter += 1
                            hash_to_index[h] = image_counter
                        img_index = hash_to_index[h]
                    break
                break
            if img_index is not None:
                blocks.append(f"[IMAGEM {img_index}]")
        elif text:
            blocks.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                blocks.append(row_text)

    return "\n".join(blocks)


def load_documents(docs_path: str) -> dict[str, str]:
    """Loads all .docx files from docs_path and returns {filename_stem: text}."""
    path = Path(docs_path)
    if not path.exists():
        raise FileNotFoundError(f"Diretório não encontrado: {docs_path}")
    return {
        file.stem: text
        for file in sorted(path.glob("*.docx"))
        if (text := _extract_text(file)).strip()
    }


def load_images(docs_path: str) -> dict[str, list[bytes]]:
    """Loads all images from .docx files and returns {filename_stem: [image_bytes]}."""
    path = Path(docs_path)
    return {
        file.stem: images
        for file in sorted(path.glob("*.docx"))
        if (images := _extract_images(file))
    }


def format_as_context(docs: dict[str, str]) -> str:
    """Formats the loaded documents into a single prompt-ready context string."""
    return "\n\n".join(f"=== {title} ===\n{content}" for title, content in docs.items())
